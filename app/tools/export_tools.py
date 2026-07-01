import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

from app.tools.workflow_tools import format_workflow_result


DEFAULT_EXPORT_DIR = "outputs/reports"


@dataclass
class ExportResult:
    """
    导出结果信息。
    """
    success: bool
    file_path: str
    file_name: str
    export_format: str
    detail: str


class ResultExportTool:
    """
    结果导出工具。

    功能：
    1. 将 workflow 输出结果导出为 Markdown；
    2. 将 workflow 输出结果导出为 Word docx；
    3. 对 Word 导出结果进行基础格式美化；
    4. 统一管理导出目录和文件命名；
    5. 为 Streamlit 下载功能提供文件生成能力。
    """

    def __init__(self, output_dir: str = DEFAULT_EXPORT_DIR):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _safe_filename(text: str) -> str:
        """
        将任务名称、标题等文本转换为安全文件名。
        """
        if not text:
            return "export_result"

        text = text.strip()
        text = re.sub(r"[\\/:*?\"<>|]", "_", text)
        text = re.sub(r"\s+", "_", text)
        text = text[:80]

        return text or "export_result"

    @staticmethod
    def _timestamp() -> str:
        """
        生成时间戳，避免文件名重复。
        """
        return datetime.now().strftime("%Y%m%d_%H%M%S")

    def _build_file_path(
        self,
        file_stem: Optional[str],
        export_format: str,
        task_type: Optional[str] = None,
    ) -> Path:
        """
        构造导出文件路径。
        """
        if file_stem:
            safe_stem = self._safe_filename(file_stem)
        elif task_type:
            safe_stem = self._safe_filename(task_type)
        else:
            safe_stem = "academic_result"

        file_name = f"{safe_stem}_{self._timestamp()}.{export_format}"
        return self.output_dir / file_name

    def export_workflow_result(
        self,
        workflow_result: Dict[str, Any],
        export_format: str = "md",
        file_stem: Optional[str] = None,
    ) -> ExportResult:
        """
        导出统一 workflow 结果。

        参数：
        - workflow_result: AcademicResearchWorkflow.run_task() 的返回结果；
        - export_format: md 或 docx；
        - file_stem: 文件名前缀，可选。
        """
        if not workflow_result:
            raise ValueError("workflow_result 不能为空")

        if export_format not in {"md", "docx"}:
            raise ValueError("export_format 只能是 md 或 docx")

        task_type = workflow_result.get("task_type", "workflow_result")
        formatted_text = format_workflow_result(workflow_result)

        title = workflow_result.get("task_name", "学术研究工作流结果")

        if export_format == "md":
            return self.export_markdown(
                content=formatted_text,
                file_stem=file_stem or task_type,
                title=title,
            )

        return self.export_docx(
            content=formatted_text,
            file_stem=file_stem or task_type,
            title=title,
        )

    def export_text(
        self,
        content: str,
        export_format: str = "md",
        file_stem: Optional[str] = None,
        title: str = "学术研究结果",
    ) -> ExportResult:
        """
        导出普通文本内容。

        适用于后续扩展：例如直接导出某个 Markdown 字符串。
        """
        if not content or not content.strip():
            raise ValueError("content 不能为空")

        if export_format == "md":
            return self.export_markdown(
                content=content,
                file_stem=file_stem,
                title=title,
            )

        if export_format == "docx":
            return self.export_docx(
                content=content,
                file_stem=file_stem,
                title=title,
            )

        raise ValueError("export_format 只能是 md 或 docx")

    def export_markdown(
        self,
        content: str,
        file_stem: Optional[str] = None,
        title: str = "学术研究结果",
    ) -> ExportResult:
        """
        导出 Markdown 文件。
        """
        file_path = self._build_file_path(
            file_stem=file_stem,
            export_format="md",
        )

        markdown_content = (
            f"# {title}\n\n"
            f"> 导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            f"---\n\n"
            f"{content}\n"
        )

        file_path.write_text(markdown_content, encoding="utf-8")

        return ExportResult(
            success=True,
            file_path=str(file_path),
            file_name=file_path.name,
            export_format="md",
            detail=f"Markdown 文件导出成功：{file_path}",
        )

    def export_docx(
        self,
        content: str,
        file_stem: Optional[str] = None,
        title: str = "学术研究结果",
    ) -> ExportResult:
        """
        导出 Word docx 文件。

        该版本对 Word 文件进行基础格式美化：
        - 设置正文、标题字体；
        - 识别 Markdown 标题；
        - 识别加粗语法 **text**；
        - 识别项目符号列表；
        - 识别编号列表；
        - 识别工作流输出中的关键小节标题。
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Inches, Pt
        except ImportError as error:
            raise ImportError(
                "未安装 python-docx。请先执行：pip install python-docx"
            ) from error

        file_path = self._build_file_path(
            file_stem=file_stem,
            export_format="docx",
        )

        document = Document()

        self._configure_docx_document(
            document=document,
            qn=qn,
            Pt=Pt,
            Inches=Inches,
        )

        title_paragraph = document.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        export_time_paragraph = document.add_paragraph()
        export_time_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        export_time_run = export_time_paragraph.add_run(
            f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        export_time_run.font.size = Pt(10)

        self._add_horizontal_rule(document, OxmlElement, qn)

        previous_blank = False

        for raw_line in content.splitlines():
            line = raw_line.rstrip()

            if not line.strip():
                if not previous_blank:
                    document.add_paragraph("")
                previous_blank = True
                continue

            previous_blank = False
            clean_line = line.strip()

            # 跳过控制台格式化分隔线
            if self._is_console_separator(clean_line):
                continue

            # Markdown 水平线
            if clean_line in {"---", "***", "___"}:
                self._add_horizontal_rule(document, OxmlElement, qn)
                continue

            # Markdown 标题
            heading_info = self._parse_markdown_heading(clean_line)
            if heading_info:
                heading_level, heading_text = heading_info
                document.add_heading(heading_text, level=heading_level)
                continue

            # 工作流格式中的中文小节标题
            if self._is_workflow_section_title(clean_line):
                document.add_heading(clean_line, level=2)
                continue

            # 【标题】格式
            if clean_line.startswith("【") and clean_line.endswith("】"):
                document.add_heading(clean_line.strip("【】"), level=2)
                continue

            # Markdown 引用
            if clean_line.startswith(">"):
                quote_text = clean_line.lstrip(">").strip()
                paragraph = document.add_paragraph(style="Intense Quote")
                self._add_runs_with_markdown_bold(paragraph, quote_text)
                continue

            # 项目符号列表
            bullet_info = self._parse_bullet_line(clean_line)
            if bullet_info:
                paragraph = document.add_paragraph(style="List Bullet")
                self._add_runs_with_markdown_bold(paragraph, bullet_info)
                continue

            # 编号列表
            numbered_info = self._parse_numbered_line(clean_line)
            if numbered_info:
                paragraph = document.add_paragraph(style="List Number")
                self._add_runs_with_markdown_bold(paragraph, numbered_info)
                continue

            # 键值行：任务类型：xxx / 任务名称：xxx 等
            if self._is_key_value_line(clean_line):
                paragraph = document.add_paragraph()
                self._add_key_value_runs(paragraph, clean_line)
                continue

            # 普通段落
            paragraph = document.add_paragraph()
            self._add_runs_with_markdown_bold(paragraph, clean_line)

        document.save(file_path)

        return ExportResult(
            success=True,
            file_path=str(file_path),
            file_name=file_path.name,
            export_format="docx",
            detail=f"Word 文件导出成功：{file_path}",
        )

    @staticmethod
    def _configure_docx_document(document, qn, Pt, Inches) -> None:
        """
        配置 Word 文档基础样式。
        """
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        styles = document.styles

        normal_style = styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.line_spacing = 1.25
        normal_style.paragraph_format.space_after = Pt(6)

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            if style_name in styles:
                style = styles[style_name]
                style.font.name = "Times New Roman"
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        if "Heading 1" in styles:
            styles["Heading 1"].font.size = Pt(16)
            styles["Heading 1"].font.bold = True

        if "Heading 2" in styles:
            styles["Heading 2"].font.size = Pt(14)
            styles["Heading 2"].font.bold = True

        if "Heading 3" in styles:
            styles["Heading 3"].font.size = Pt(12)
            styles["Heading 3"].font.bold = True

    @staticmethod
    def _is_console_separator(line: str) -> bool:
        """
        判断是否为控制台输出中的分隔线。
        """
        if not line:
            return False

        return (
            len(line) >= 5
            and set(line) <= {"=", "-", "_"}
        )

    @staticmethod
    def _parse_markdown_heading(line: str) -> Optional[tuple[int, str]]:
        """
        解析 Markdown 标题。
        """
        match = re.match(r"^(#{1,6})\s+(.+)$", line)

        if not match:
            return None

        hashes, text = match.groups()
        level = min(len(hashes), 3)

        return level, text.strip()

    @staticmethod
    def _is_workflow_section_title(line: str) -> bool:
        """
        识别 format_workflow_result 输出中的关键小节标题。
        """
        section_titles = {
            "统一学术研究工作流结果",
            "任务输出",
            "问题",
            "回答",
            "依据来源",
            "不确定之处",
            "精读结果",
            "对比结果",
            "综合对比分析",
            "文献综述框架",
            "检查结果",
        }

        return line in section_titles

    @staticmethod
    def _parse_bullet_line(line: str) -> Optional[str]:
        """
        解析项目符号列表。
        """
        match = re.match(r"^[-*+]\s+(.+)$", line)

        if not match:
            return None

        return match.group(1).strip()

    @staticmethod
    def _parse_numbered_line(line: str) -> Optional[str]:
        """
        解析编号列表。

        支持：
        1. 内容
        1、内容
        （1）内容
        """
        patterns = [
            r"^\d+[.]\s+(.+)$",
            r"^\d+[、]\s*(.+)$",
            r"^（\d+）\s*(.+)$",
            r"^\(\d+\)\s*(.+)$",
        ]

        for pattern in patterns:
            match = re.match(pattern, line)
            if match:
                return match.group(1).strip()

        return None

    @staticmethod
    def _is_key_value_line(line: str) -> bool:
        """
        识别常见键值行。
        """
        prefixes = [
            "任务类型：",
            "任务名称：",
            "任务状态：",
            "论文文件：",
            "精读维度数量：",
            "论文数量：",
            "对比维度数量：",
            "来源类型=",
            "摘要维度=",
        ]

        return any(line.startswith(prefix) for prefix in prefixes)

    @staticmethod
    def _add_key_value_runs(paragraph, line: str) -> None:
        """
        将“键：值”形式写成“键”加粗、“值”正常。
        """
        if "：" in line:
            key, value = line.split("：", 1)
            key_run = paragraph.add_run(f"{key}：")
            key_run.bold = True
            paragraph.add_run(value.strip())
            return

        if "=" in line:
            key, value = line.split("=", 1)
            key_run = paragraph.add_run(f"{key}=")
            key_run.bold = True
            paragraph.add_run(value.strip())
            return

        paragraph.add_run(line)

    @staticmethod
    def _add_runs_with_markdown_bold(paragraph, text: str) -> None:
        """
        将包含 **加粗** 的文本写入段落。

        示例：
        人工智能通过 **信息处理效率** 影响企业创新韧性。
        """
        if not text:
            return

        parts = re.split(r"(\*\*.*?\*\*)", text)

        for part in parts:
            if not part:
                continue

            if part.startswith("**") and part.endswith("**") and len(part) >= 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    @staticmethod
    def _add_horizontal_rule(document, OxmlElement, qn) -> None:
        """
        添加水平分隔线。
        """
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 6

        p = paragraph._p
        p_pr = p.get_or_add_pPr()

        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")

        p_bdr.append(bottom)
        p_pr.append(p_bdr)


def format_export_result(export_result: ExportResult) -> str:
    """
    将导出结果格式化为终端可读文本。
    """
    lines = []

    lines.append("=" * 80)
    lines.append("结果导出完成")
    lines.append("=" * 80)
    lines.append(f"导出状态：{'成功' if export_result.success else '失败'}")
    lines.append(f"导出格式：{export_result.export_format}")
    lines.append(f"文件名称：{export_result.file_name}")
    lines.append(f"文件路径：{export_result.file_path}")
    lines.append(f"说明：{export_result.detail}")

    return "\n".join(lines)